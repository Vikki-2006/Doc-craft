"""
PDF Toolkit Pro - Core Engine (pdf_utils.py)
Implements all PDF manipulation algorithms using PyMuPDF (fitz), pypdf, Pillow, and ReportLab.
100% offline, efficient, and robust.
"""

import os
import io
import fitz  # PyMuPDF
import pypdf
from PIL import Image

from reportlab.pdfgen import canvas
from reportlab.lib.colors import HexColor

from .helpers import (
    PDFToolkitError,
    PasswordProtectedPDFError,
    CorruptedPDFError,
    validate_page_range
)

def get_pdf_info(pdf_path):
    """
    Inspect a PDF file and return metadata dict:
    - total_pages: int
    - file_size: int
    - is_encrypted: bool
    - title: str
    - author: str
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"File not found: {pdf_path}")

    file_size = os.path.getsize(pdf_path)
    
    try:
        doc = fitz.open(pdf_path)
        if doc.is_encrypted:
            raise PasswordProtectedPDFError(f"PDF file is password protected: {os.path.basename(pdf_path)}")
            
        metadata = doc.metadata or {}
        total_pages = doc.page_count
        doc.close()
        
        return {
            "file_path": pdf_path,
            "filename": os.path.basename(pdf_path),
            "total_pages": total_pages,
            "file_size": file_size,
            "is_encrypted": False,
            "title": metadata.get("title", ""),
            "author": metadata.get("author", "")
        }
    except PasswordProtectedPDFError:
        raise
    except Exception as e:
        raise CorruptedPDFError(f"Failed to read PDF file '{os.path.basename(pdf_path)}': {str(e)}")


def render_page_thumbnail(pdf_path, page_num=0, max_size=(250, 350), dpi=96):
    """
    Render a single PDF page into a PIL Image for UI display.
    """
    try:
        doc = fitz.open(pdf_path)
        if doc.is_encrypted:
            raise PasswordProtectedPDFError("PDF is password protected.")
        if page_num < 0 or page_num >= doc.page_count:
            page_num = 0
            
        page = doc.load_page(page_num)
        pix = page.get_pixmap(dpi=dpi)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        doc.close()
        
        img.thumbnail(max_size, Image.Resampling.LANCZOS)
        return img
    except Exception as e:
        # Fallback error thumbnail image
        err_img = Image.new("RGB", max_size, (239, 68, 68))
        return err_img


def merge_pdfs(pdf_paths, output_path, progress_callback=None):
    """
    Merge multiple PDF files into a single PDF file.
    - pdf_paths: list of PDF file paths in desired order
    - output_path: output PDF path
    - progress_callback: func(current_file_idx, total_files, status_msg)
    """
    if not pdf_paths or len(pdf_paths) < 2:
        raise ValueError("At least 2 PDF files are required to perform a merge operation.")

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    
    merged_doc = fitz.open()
    total_files = len(pdf_paths)

    try:
        for idx, pdf_path in enumerate(pdf_paths):
            if progress_callback:
                progress_callback(idx, total_files, f"Merging file {idx + 1}/{total_files}: {os.path.basename(pdf_path)}")
                
            if not os.path.exists(pdf_path):
                raise FileNotFoundError(f"Source file not found: {pdf_path}")
                
            src_doc = fitz.open(pdf_path)
            if src_doc.is_encrypted:
                src_doc.close()
                raise PasswordProtectedPDFError(f"File '{os.path.basename(pdf_path)}' is password protected.")
                
            merged_doc.insert_pdf(src_doc)
            src_doc.close()

        if progress_callback:
            progress_callback(total_files - 1, total_files, "Saving merged PDF document...")

        merged_doc.save(output_path, garbage=4, deflate=True)
        merged_doc.close()

        if progress_callback:
            progress_callback(total_files, total_files, "Merge completed successfully!")

        return output_path
    except Exception as e:
        merged_doc.close()
        if isinstance(e, PDFToolkitError):
            raise e
        raise PDFToolkitError(f"Merge operation failed: {str(e)}")


def split_pdf(pdf_path, output_dir, mode="every", range_str=None, chunk_size=1, progress_callback=None):
    """
    Split a PDF into multiple output files or range files.
    - mode: 'every' (individual page per PDF), 'range' (specific page selection into 1 output file), 'chunk' (N pages per file)
    - range_str: syntax like '1-3, 5, 7-10' when mode == 'range'
    - chunk_size: integer for 'chunk' mode
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"Source file not found: {pdf_path}")

    os.makedirs(output_dir, exist_ok=True)
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    
    doc = fitz.open(pdf_path)
    if doc.is_encrypted:
        doc.close()
        raise PasswordProtectedPDFError(f"PDF file is password protected: {os.path.basename(pdf_path)}")

    total_pages = doc.page_count
    created_files = []

    try:
        if mode == "every":
            for page_idx in range(total_pages):
                if progress_callback:
                    progress_callback(page_idx, total_pages, f"Extracting page {page_idx + 1}/{total_pages}...")
                
                out_doc = fitz.open()
                out_doc.insert_pdf(doc, from_page=page_idx, to_page=page_idx)
                
                out_path = os.path.join(output_dir, f"{base_name}_page_{page_idx + 1:03d}.pdf")
                out_doc.save(out_path, garbage=4, deflate=True)
                out_doc.close()
                created_files.append(out_path)

        elif mode == "range":
            if not range_str:
                raise ValueError("Range string is required for range split mode.")
            selected_indices = validate_page_range(range_str, total_pages)
            
            if progress_callback:
                progress_callback(0, len(selected_indices), "Extracting requested page range...")
                
            out_doc = fitz.open()
            for idx in selected_indices:
                out_doc.insert_pdf(doc, from_page=idx, to_page=idx)
                
            out_path = os.path.join(output_dir, f"{base_name}_extracted_range.pdf")
            out_doc.save(out_path, garbage=4, deflate=True)
            out_doc.close()
            created_files.append(out_path)

        elif mode == "chunk":
            chunk_size = max(1, int(chunk_size))
            chunk_count = (total_pages + chunk_size - 1) // chunk_size
            
            for chunk_idx in range(chunk_count):
                start_p = chunk_idx * chunk_size
                end_p = min(start_p + chunk_size - 1, total_pages - 1)
                
                if progress_callback:
                    progress_callback(chunk_idx, chunk_count, f"Creating split part {chunk_idx + 1}/{chunk_count}...")
                    
                out_doc = fitz.open()
                out_doc.insert_pdf(doc, from_page=start_p, to_page=end_p)
                
                out_path = os.path.join(output_dir, f"{base_name}_part_{chunk_idx + 1:02d}_pages_{start_p+1}-{end_p+1}.pdf")
                out_doc.save(out_path, garbage=4, deflate=True)
                out_doc.close()
                created_files.append(out_path)

        doc.close()
        
        if progress_callback:
            progress_callback(100, 100, f"Successfully split into {len(created_files)} PDF file(s)!")
            
        return created_files

    except Exception as e:
        doc.close()
        if isinstance(e, PDFToolkitError):
            raise e
        raise PDFToolkitError(f"Split operation failed: {str(e)}")


def rotate_pdf(pdf_path, output_path, rotation_angle=90, page_selection="all", range_str=None, progress_callback=None):
    """
    Rotate pages in a PDF document.
    - rotation_angle: 90 (CW), 180, 270 (CCW)
    - page_selection: 'all', 'odd', 'even', 'range'
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"Source file not found: {pdf_path}")

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    
    doc = fitz.open(pdf_path)
    if doc.is_encrypted:
        doc.close()
        raise PasswordProtectedPDFError("PDF is password protected.")

    total_pages = doc.page_count
    target_indices = []

    if page_selection == "all":
        target_indices = list(range(total_pages))
    elif page_selection == "odd":
        target_indices = [i for i in range(total_pages) if (i + 1) % 2 != 0]
    elif page_selection == "even":
        target_indices = [i for i in range(total_pages) if (i + 1) % 2 == 0]
    elif page_selection == "range":
        target_indices = validate_page_range(range_str, total_pages)

    try:
        for idx, page_idx in enumerate(target_indices):
            if progress_callback:
                progress_callback(idx, len(target_indices), f"Rotating page {page_idx + 1}/{total_pages}...")
                
            page = doc.load_page(page_idx)
            current_rot = page.rotation
            page.set_rotation((current_rot + rotation_angle) % 360)

        if progress_callback:
            progress_callback(len(target_indices), len(target_indices), "Saving rotated PDF...")

        doc.save(output_path, garbage=4, deflate=True)
        doc.close()
        return output_path

    except Exception as e:
        doc.close()
        if isinstance(e, PDFToolkitError):
            raise e
        raise PDFToolkitError(f"Rotation failed: {str(e)}")


def pdf_to_images(pdf_path, output_dir, fmt="png", dpi=150, range_str=None, quality=90, progress_callback=None):
    """
    Export PDF pages as raster images (PNG or JPEG) at specified DPI resolution.
    - fmt: 'png' or 'jpeg'
    - dpi: 72, 150, or 300
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"Source file not found: {pdf_path}")

    os.makedirs(output_dir, exist_ok=True)
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    
    doc = fitz.open(pdf_path)
    if doc.is_encrypted:
        doc.close()
        raise PasswordProtectedPDFError("PDF file is password protected.")

    total_pages = doc.page_count
    target_indices = validate_page_range(range_str, total_pages) if range_str else list(range(total_pages))
    
    created_images = []
    fmt = fmt.lower()
    if fmt == "jpg":
        fmt = "jpeg"

    try:
        for idx, page_idx in enumerate(target_indices):
            if progress_callback:
                progress_callback(idx, len(target_indices), f"Exporting page {page_idx + 1}/{total_pages} at {dpi} DPI...")
                
            page = doc.load_page(page_idx)
            pix = page.get_pixmap(dpi=dpi)
            
            ext = "jpg" if fmt == "jpeg" else "png"
            out_path = os.path.join(output_dir, f"{base_name}_page_{page_idx + 1:03d}.{ext}")
            
            if fmt == "jpeg":
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                img.save(out_path, "JPEG", quality=quality)
            else:
                pix.save(out_path)
                
            created_images.append(out_path)

        doc.close()
        
        if progress_callback:
            progress_callback(len(target_indices), len(target_indices), f"Successfully exported {len(created_images)} image(s)!")
            
        return created_images

    except Exception as e:
        doc.close()
        if isinstance(e, PDFToolkitError):
            raise e
        raise PDFToolkitError(f"PDF to image export failed: {str(e)}")


def add_text_watermark(pdf_path, output_path, text, opacity=0.3, position="center", font_size=40, angle=45, color_hex="#3B82F6", range_str=None, progress_callback=None):
    """
    Apply a custom text watermark onto PDF pages using ReportLab for precision rendering.
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"Source file not found: {pdf_path}")

    if not text or not text.strip():
        raise ValueError("Watermark text cannot be empty.")

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    
    doc = fitz.open(pdf_path)
    if doc.is_encrypted:
        doc.close()
        raise PasswordProtectedPDFError("PDF file is password protected.")

    total_pages = doc.page_count
    target_indices = validate_page_range(range_str, total_pages) if range_str else list(range(total_pages))

    try:
        for idx, page_idx in enumerate(target_indices):
            if progress_callback:
                progress_callback(idx, len(target_indices), f"Applying text watermark to page {page_idx + 1}/{total_pages}...")
                
            page = doc.load_page(page_idx)
            rect = page.rect
            width, height = rect.width, rect.height
            
            # Generate ReportLab overlay PDF in memory
            packet = io.BytesIO()
            can = canvas.Canvas(packet, pagesize=(width, height))
            can.setFillColor(HexColor(color_hex), alpha=opacity)
            can.setFont("Helvetica-Bold", font_size)
            
            can.saveState()
            
            # Compute position
            if position == "center":
                cx, cy = width / 2, height / 2
                can.translate(cx, cy)
                can.rotate(angle)
                can.drawCentredString(0, 0, text)
            elif position == "top_left":
                can.translate(60, height - 60)
                can.rotate(angle)
                can.drawString(0, 0, text)
            elif position == "top_right":
                can.translate(width - 60, height - 60)
                can.rotate(angle)
                can.drawRightString(0, 0, text)
            elif position == "bottom_left":
                can.translate(60, 60)
                can.rotate(angle)
                can.drawString(0, 0, text)
            elif position == "bottom_right":
                can.translate(width - 60, 60)
                can.rotate(angle)
                can.drawRightString(0, 0, text)
            elif position == "tile":
                # Grid pattern tile
                for x in range(50, int(width), 180):
                    for y in range(50, int(height), 150):
                        can.saveState()
                        can.translate(x, y)
                        can.rotate(angle)
                        can.drawCentredString(0, 0, text)
                        can.restoreState()
                        
            can.restoreState()
            can.save()
            
            packet.seek(0)
            overlay_pdf = fitz.open("pdf", packet.read())
            page.show_pdf_page(rect, overlay_pdf, 0)
            overlay_pdf.close()

        if progress_callback:
            progress_callback(len(target_indices), len(target_indices), "Saving watermarked PDF...")

        doc.save(output_path, garbage=4, deflate=True)
        doc.close()
        return output_path

    except Exception as e:
        doc.close()
        if isinstance(e, PDFToolkitError):
            raise e
        raise PDFToolkitError(f"Text watermark operation failed: {str(e)}")


def add_image_watermark(pdf_path, output_path, watermark_img_path, opacity=0.3, position="center", scale=0.5, range_str=None, progress_callback=None):
    """
    Apply an image logo watermark onto PDF pages.
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"Source file not found: {pdf_path}")
    if not os.path.exists(watermark_img_path):
        raise FileNotFoundError(f"Watermark image not found: {watermark_img_path}")

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    
    doc = fitz.open(pdf_path)
    if doc.is_encrypted:
        doc.close()
        raise PasswordProtectedPDFError("PDF file is password protected.")

    total_pages = doc.page_count
    target_indices = validate_page_range(range_str, total_pages) if range_str else list(range(total_pages))

    try:
        # Pre-process watermark image with opacity
        wm_img = Image.open(watermark_img_path).convert("RGBA")
        r, g, b, a = wm_img.split()
        a = a.point(lambda p: int(p * opacity))
        wm_img.putalpha(a)
        
        # Save temp transparent watermark image
        temp_wm_path = os.path.join(os.path.dirname(output_path), "_temp_wm.png")
        wm_img.save(temp_wm_path, "PNG")

        for idx, page_idx in enumerate(target_indices):
            if progress_callback:
                progress_callback(idx, len(target_indices), f"Applying image watermark to page {page_idx + 1}/{total_pages}...")
                
            page = doc.load_page(page_idx)
            rect = page.rect
            pw, ph = rect.width, rect.height
            
            # Calculate watermark target dimensions
            w_w = pw * scale
            w_h = (wm_img.height / wm_img.width) * w_w
            
            if position == "center":
                x0 = (pw - w_w) / 2
                y0 = (ph - w_h) / 2
            elif position == "top_left":
                x0 = 30
                y0 = 30
            elif position == "top_right":
                x0 = pw - w_w - 30
                y0 = 30
            elif position == "bottom_left":
                x0 = 30
                y0 = ph - w_h - 30
            elif position == "bottom_right":
                x0 = pw - w_w - 30
                y0 = ph - w_h - 30
            else:
                x0 = (pw - w_w) / 2
                y0 = (ph - w_h) / 2

            wm_rect = fitz.Rect(x0, y0, x0 + w_w, y0 + w_h)
            page.insert_image(wm_rect, filename=temp_wm_path, overlay=True)

        if progress_callback:
            progress_callback(len(target_indices), len(target_indices), "Saving watermarked PDF...")

        doc.save(output_path, garbage=4, deflate=True)
        doc.close()
        
        if os.path.exists(temp_wm_path):
            os.remove(temp_wm_path)
            
        return output_path

    except Exception as e:
        doc.close()
        if isinstance(e, PDFToolkitError):
            raise e
        raise PDFToolkitError(f"Image watermark operation failed: {str(e)}")


# ─────────────────────────────────────────────────────────────────────────────
# NEW TOOLS — DocCraft 2.0
# ─────────────────────────────────────────────────────────────────────────────

def compress_pdf(pdf_path, output_path, progress_callback=None):
    """
    Compress a PDF by removing unused objects and optimizing streams/images/fonts.
    Returns (output_path, original_size, compressed_size, reduction_percent).
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"Source file not found: {pdf_path}")

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

    doc = fitz.open(pdf_path)
    if doc.is_encrypted:
        doc.close()
        raise PasswordProtectedPDFError("PDF file is password protected.")

    original_size = os.path.getsize(pdf_path)

    try:
        if progress_callback:
            progress_callback(0, 1, "Optimising PDF structure and streams…")

        save_opts = dict(garbage=4, deflate=True, clean=True)
        try:
            doc.save(output_path, deflate_images=True, deflate_fonts=True, **save_opts)
        except TypeError:
            # Older PyMuPDF build — fall back to basic compression
            doc.save(output_path, **save_opts)

        doc.close()

        compressed_size = os.path.getsize(output_path)
        reduction = max(0.0, round((1 - compressed_size / original_size) * 100, 1))

        if progress_callback:
            progress_callback(1, 1, f"Compression complete — {reduction}% size reduction.")

        return output_path, original_size, compressed_size, reduction

    except Exception as e:
        doc.close()
        if isinstance(e, PDFToolkitError):
            raise
        raise PDFToolkitError(f"Compression failed: {str(e)}")


def add_password(pdf_path, output_path, user_password, owner_password=None, progress_callback=None):
    """
    Protect a PDF with AES-256 encryption using a user password.
    owner_password defaults to user_password if not provided.
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"Source file not found: {pdf_path}")
    if not user_password:
        raise ValueError("Password cannot be empty.")

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

    doc = fitz.open(pdf_path)
    if doc.is_encrypted:
        doc.close()
        raise PasswordProtectedPDFError("PDF is already password protected. Remove the existing password first.")

    if not owner_password:
        owner_password = user_password

    try:
        if progress_callback:
            progress_callback(0, 1, "Applying AES-256 encryption…")

        perm = int(
            fitz.PDF_PERM_ACCESSIBILITY
            | fitz.PDF_PERM_PRINT
            | fitz.PDF_PERM_COPY
        )

        doc.save(
            output_path,
            encryption=fitz.PDF_ENCRYPT_AES_256,
            owner_pw=owner_password,
            user_pw=user_password,
            permissions=perm,
            garbage=4,
            deflate=True,
        )
        doc.close()

        if progress_callback:
            progress_callback(1, 1, "Password protection applied.")

        return output_path

    except Exception as e:
        doc.close()
        if isinstance(e, PDFToolkitError):
            raise
        raise PDFToolkitError(f"Password encryption failed: {str(e)}")


def remove_password(pdf_path, output_path, password, progress_callback=None):
    """
    Remove password protection from a PDF by re-saving without encryption.
    Raises ValueError if the password is incorrect.
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"Source file not found: {pdf_path}")

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

    doc = fitz.open(pdf_path)
    if doc.is_encrypted:
        if not doc.authenticate(password):
            doc.close()
            raise ValueError("Incorrect password — cannot unlock this PDF.")

    try:
        if progress_callback:
            progress_callback(0, 1, "Removing encryption…")

        doc.save(
            output_path,
            encryption=fitz.PDF_ENCRYPT_NONE,
            garbage=4,
            deflate=True,
        )
        doc.close()

        if progress_callback:
            progress_callback(1, 1, "Password removed successfully.")

        return output_path

    except Exception as e:
        doc.close()
        if isinstance(e, PDFToolkitError):
            raise
        raise PDFToolkitError(f"Remove password failed: {str(e)}")


def add_page_numbers(
    pdf_path,
    output_path,
    position="bottom-center",
    start_num=1,
    prefix="",
    suffix="",
    font_size=11,
    progress_callback=None,
):
    """
    Insert page number text at a specified position on every page.
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"Source file not found: {pdf_path}")

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

    doc = fitz.open(pdf_path)
    if doc.is_encrypted:
        doc.close()
        raise PasswordProtectedPDFError("PDF file is password protected.")

    total_pages = doc.page_count
    margin = 26

    try:
        for page_idx, page in enumerate(doc):
            if progress_callback:
                progress_callback(page_idx, total_pages, f"Numbering page {page_idx + 1}/{total_pages}…")

            rect = page.rect
            label = f"{prefix}{page_idx + start_num}{suffix}"

            # Determine insertion point
            if position == "bottom-center":
                pt = fitz.Point(rect.width / 2, rect.height - margin)
            elif position == "bottom-left":
                pt = fitz.Point(margin, rect.height - margin)
            elif position == "bottom-right":
                pt = fitz.Point(rect.width - margin, rect.height - margin)
            elif position == "top-center":
                pt = fitz.Point(rect.width / 2, margin + font_size)
            elif position == "top-left":
                pt = fitz.Point(margin, margin + font_size)
            elif position == "top-right":
                pt = fitz.Point(rect.width - margin, margin + font_size)
            else:
                pt = fitz.Point(rect.width / 2, rect.height - margin)

            page.insert_text(
                pt,
                label,
                fontsize=font_size,
                color=(0.3, 0.3, 0.3),
                fontname="Helvetica",
            )

        if progress_callback:
            progress_callback(total_pages, total_pages, "Page numbers added.")

        doc.save(output_path, garbage=4, deflate=True)
        doc.close()
        return output_path

    except Exception as e:
        doc.close()
        if isinstance(e, PDFToolkitError):
            raise
        raise PDFToolkitError(f"Page numbering failed: {str(e)}")


def pdf_to_word(pdf_path, output_path, progress_callback=None):
    """
    Export PDF text content to a Microsoft Word (.docx) document.
    Extracts text per page; layout/images from scanned PDFs are not reconstructed.
    """
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Cm
    except ImportError:
        raise PDFToolkitError(
            "python-docx is required. Install it with: pip install python-docx"
        )

    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"Source file not found: {pdf_path}")

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

    doc_in = fitz.open(pdf_path)
    if doc_in.is_encrypted:
        doc_in.close()
        raise PasswordProtectedPDFError("PDF file is password protected.")

    total_pages = doc_in.page_count
    docx_out = DocxDocument()

    # Set reasonable margins
    section = docx_out.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    try:
        for page_idx in range(total_pages):
            if progress_callback:
                progress_callback(page_idx, total_pages, f"Extracting page {page_idx + 1}/{total_pages}…")

            page = doc_in.load_page(page_idx)
            text = page.get_text("text")

            if page_idx > 0:
                docx_out.add_page_break()

            h = docx_out.add_heading(f"Page {page_idx + 1}", level=2)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(0x43, 0x61, 0xEE)

            if text.strip():
                for line in text.split("\n"):
                    stripped = line.strip()
                    if stripped:
                        p = docx_out.add_paragraph(stripped)
                        if p.runs:
                            p.runs[0].font.size = Pt(10)
            else:
                p = docx_out.add_paragraph("[No extractable text on this page]")
                if p.runs:
                    p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        if progress_callback:
            progress_callback(total_pages, total_pages, "Word document saved.")

        docx_out.save(output_path)
        doc_in.close()
        return output_path

    except Exception as e:
        doc_in.close()
        if isinstance(e, PDFToolkitError):
            raise
        raise PDFToolkitError(f"PDF to Word conversion failed: {str(e)}")


def pdf_to_excel(pdf_path, output_path, progress_callback=None):
    """
    Export PDF text content to a Microsoft Excel (.xlsx) workbook.
    Creates a Summary sheet and one sheet per PDF page.
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        raise PDFToolkitError(
            "openpyxl is required. Install it with: pip install openpyxl"
        )

    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"Source file not found: {pdf_path}")

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

    doc_in = fitz.open(pdf_path)
    if doc_in.is_encrypted:
        doc_in.close()
        raise PasswordProtectedPDFError("PDF file is password protected.")

    total_pages = doc_in.page_count
    wb = Workbook()

    # Summary sheet
    ws_summary = wb.active
    ws_summary.title = "Summary"
    bold_font = Font(bold=True)
    ws_summary.append(["DocCraft — PDF to Excel Export"])
    ws_summary.append(["Source", os.path.basename(pdf_path)])
    ws_summary.append(["Pages", total_pages])
    ws_summary.append([])
    ws_summary.append(["Page", "Text Lines", "Characters"])
    for cell in ws_summary["5"]:
        cell.font = bold_font

    try:
        for page_idx in range(total_pages):
            if progress_callback:
                progress_callback(page_idx, total_pages, f"Processing page {page_idx + 1}/{total_pages}…")

            page = doc_in.load_page(page_idx)
            text = page.get_text("text")
            lines = [l.strip() for l in text.split("\n") if l.strip()]

            ws_summary.append([page_idx + 1, len(lines), len(text)])

            sheet_name = f"Page {page_idx + 1}"[:31]
            ws = wb.create_sheet(title=sheet_name)
            ws.append([f"Page {page_idx + 1} — Extracted Text"])
            ws.append(["Line #", "Content"])
            ws["A1"].font = bold_font
            ws["A2"].font = bold_font
            ws["B2"].font = bold_font

            if lines:
                for i, line in enumerate(lines, 1):
                    ws.append([i, line])
            else:
                ws.append(["—", "[No extractable text on this page]"])

            ws.column_dimensions["A"].width = 8
            ws.column_dimensions["B"].width = 90

        if progress_callback:
            progress_callback(total_pages, total_pages, "Excel workbook saved.")

        ws_summary.column_dimensions["A"].width = 14
        ws_summary.column_dimensions["B"].width = 50
        ws_summary.column_dimensions["C"].width = 14

        wb.save(output_path)
        doc_in.close()
        return output_path

    except Exception as e:
        doc_in.close()
        if isinstance(e, PDFToolkitError):
            raise
        raise PDFToolkitError(f"PDF to Excel conversion failed: {str(e)}")
